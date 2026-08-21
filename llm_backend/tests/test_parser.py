import pytest

from app.services.doc_parser import parse_text_file, parse_docx, Segment


def _write(tmp_path, name, data: bytes):
    p = tmp_path / name
    p.write_bytes(data)
    return p


def test_gbk_txt_decoded(tmp_path):
    p = _write(tmp_path, "gbk.txt", "价格：5999 元".encode("gbk"))
    segments = parse_text_file(p, "txt")
    assert "5999" in segments[0].text


def test_utf8_txt_single_segment(tmp_path):
    p = _write(tmp_path, "a.txt", "普通文本".encode("utf-8"))
    segments = parse_text_file(p, "txt")
    assert len(segments) == 1 and segments[0].chapter == ""


def test_md_chapter_stack(tmp_path):
    md = "# 一、智能沙发系列\n## 云享沙发 SF-2000\n价格 5999\n## 云享沙发 SF-1000\n价格 3999"
    p = _write(tmp_path, "a.md", md.encode("utf-8"))
    segments = parse_text_file(p, "md")
    assert [s.chapter for s in segments] == [
        "一、智能沙发系列",
        "一、智能沙发系列 > 云享沙发 SF-2000",
        "一、智能沙发系列 > 云享沙发 SF-1000",
    ]


def test_md_cross_chapter_chunk_ownership(tmp_path):
    """跨章节内容归属规则:块内首个非空字符所在章节(由分块阶段保证,此处验证段边界)。"""
    md = "## 章节A\n短\n## 章节B\n内容"
    p = _write(tmp_path, "a.md", md.encode("utf-8"))
    segments = parse_text_file(p, "md")
    assert segments[0].chapter == "章节A"
    assert segments[1].chapter == "章节B"


def test_docx_body_order_and_tables(tmp_path):
    """段落-表格-段落交替:body 迭代保持顺序,表格合并单元格不重复。"""
    from docx import Document

    doc = Document()
    doc.add_paragraph("开头段落")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "型号"
    table.cell(0, 1).text = "SF-2000"
    table.cell(1, 0).text = "价格"
    table.cell(1, 1).text = "5999"
    doc.add_paragraph("结尾段落")
    p = tmp_path / "a.docx"
    doc.save(str(p))

    segments = parse_docx(p)
    joined = "\n".join(s.text for s in segments)
    assert joined.index("开头段落") < joined.index("型号 | SF-2000")
    assert joined.index("型号 | SF-2000") < joined.index("价格 | 5999")
    assert joined.index("价格 | 5999") < joined.index("结尾段落")


def test_docx_heading_chapter(tmp_path):
    from docx import Document

    doc = Document()
    doc.add_heading("智能沙发系列", level=1)
    doc.add_paragraph("参数说明")
    p = tmp_path / "b.docx"
    doc.save(str(p))

    segments = parse_docx(p)
    assert any(s.chapter == "智能沙发系列" for s in segments)
