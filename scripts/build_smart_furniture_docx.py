"""
将 scripts/data/jd_smart_furniture.tsv（真实京东在售智能家具商品清单）转换为
单份层级 docx 商品知识文档，供手动上传索引构建 / 检索测试。

文档结构（可被索引管道 docx 解析按章节切分）:
  H1 京东智能家具产品知识文档
  ├─ 文档说明
  ├─ H2 一、电动智能沙发（品类, 中文序号）
  │   └─ H3 商品名称
  │       ├─ H4 商品信息（品牌/品类/参考价格）
  │       ├─ H4 功能特点（要点列表）
  │       ├─ H4 规格参数（要点列表）
  │       ├─ H4 售后服务（要点列表, 有数据才生成）
  │       └─ H4 参考来源
  ...

用法:
  python scripts/build_smart_furniture_docx.py [--limit 50]

输出: llm_backend/knowledge_data/product_knowledge_docx/京东智能家具产品知识文档.docx (gitignore 内)
"""

import argparse
import csv
import sys
from pathlib import Path

from docx import Document

PROJECT_ROOT = Path(__file__).resolve().parent.parent
TSV_PATH = PROJECT_ROOT / "scripts" / "data" / "jd_smart_furniture.tsv"
OUTPUT_DIR = PROJECT_ROOT / "llm_backend" / "knowledge_data" / "product_knowledge_docx"
DOC_TITLE = "京东智能家具产品知识文档"

# TSV 列名
COL_CATEGORY, COL_NAME, COL_BRAND, COL_PRICE = "品类", "商品名称", "品牌", "参考价格(元)"
COL_FEATURES, COL_SPECS, COL_AFTERSALES, COL_SOURCE = "功能特点", "规格参数", "售后服务", "来源"

# 品类中文序号
CN_NUM = "一二三四五六七八九十"


def split_bullets(text: str) -> list[str]:
    """按分号切分为要点列表。"""
    return [t.strip() for t in text.split("；") if t.strip()]


def add_product(doc: Document, row: dict) -> None:
    """单个商品: H3 商品名称 + H4 子章节。"""
    doc.add_heading(row[COL_NAME].strip(), level=3)

    # 商品信息（价格为动态信息，存储于数据库 product_price_stock 表，不入文档）
    doc.add_heading("商品信息", level=4)
    if row[COL_BRAND].strip():
        doc.add_paragraph(f"品牌：{row[COL_BRAND].strip()}")
    if row[COL_CATEGORY].strip():
        doc.add_paragraph(f"品类：{row[COL_CATEGORY].strip()}")

    # 功能特点
    if row[COL_FEATURES].strip():
        doc.add_heading("功能特点", level=4)
        for item in split_bullets(row[COL_FEATURES]):
            doc.add_paragraph(item, style="List Bullet")

    # 规格参数
    if row[COL_SPECS].strip():
        doc.add_heading("规格参数", level=4)
        for item in split_bullets(row[COL_SPECS]):
            doc.add_paragraph(item, style="List Bullet")

    # 售后服务（京东自营 → 引用独立政策文档；否则原文要点）
    aftersales = row[COL_AFTERSALES].strip()
    if aftersales == "京东自营":
        doc.add_heading("售后服务", level=4)
        doc.add_paragraph(
            "本商品为京东自营，适用《京东自营售后政策》（独立知识文档，已随知识库上传）",
            style="List Bullet",
        )
    elif aftersales:
        doc.add_heading("售后服务", level=4)
        for item in split_bullets(aftersales):
            doc.add_paragraph(item, style="List Bullet")

    # 参考来源
    if row[COL_SOURCE].strip():
        doc.add_heading("参考来源", level=4)
        doc.add_paragraph(row[COL_SOURCE].strip())


def main():
    parser = argparse.ArgumentParser(description="京东智能家具商品清单 → 单份层级 docx 知识文档")
    parser.add_argument("--limit", type=int, default=0, help="最多收录多少款商品（默认全部）")
    args = parser.parse_args()

    with open(TSV_PATH, encoding="utf-8") as f:
        rows = list(csv.DictReader(f, delimiter="\t"))
    if args.limit:
        rows = rows[: args.limit]
    print(f"清单共 {len(rows)} 款商品")

    doc = Document()
    doc.add_heading(DOC_TITLE, level=1)
    doc.add_paragraph(
        "本文档涵盖京东在售智能家具商品的真实信息，包括电动智能沙发、按摩椅、智能电动床、"
        "智能床垫、电动升降桌、智能晾衣架、智能窗帘、智能门锁、智能床头柜九大品类，"
        "共计 50 款商品。商品名称、品牌、型号、参数均来自京东及公开渠道的真实在售信息，"
        "整理日期：2026年8月。"
    )
    doc.add_paragraph(
        "商品价格与库存为动态信息，存储于系统数据库中（product_price_stock 表），本文档不包含。"
    )

    # 按品类分组（保持 TSV 出现顺序），中文序号
    seen_categories, category_products = [], {}
    for row in rows:
        cat = row[COL_CATEGORY].strip()
        if cat not in category_products:
            category_products[cat] = []
            seen_categories.append(cat)
        category_products[cat].append(row)

    for idx, cat in enumerate(seen_categories, start=1):
        num = CN_NUM[idx - 1] if idx <= len(CN_NUM) else str(idx)
        doc.add_heading(f"{num}、{cat}", level=2)
        for row in category_products[cat]:
            add_product(doc, row)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    out_path = OUTPUT_DIR / f"{DOC_TITLE}.docx"
    doc.save(str(out_path))
    print(f"完成: 生成 {out_path}（{len(rows)} 款商品, {len(seen_categories)} 个品类）")


if __name__ == "__main__":
    main()
