"""
将 jd-aig CEPSUM 数据集（京东多模态商品摘要，NLPCC 2022）转换为 docx 商品知识文档，
供手动上传索引构建 / 检索测试使用。

数据格式（以官方发放为准，本脚本自动探测）：
  每条记录包含: 商品描述(Input text) + 属性知识库(Product attribute) + 商品摘要(Output summary)
  常见文件格式: jsonl / json / csv / txt

用法:
  python scripts/generate_docx_from_jdaig.py --data-dir <数据集目录> [--limit 50] [--keywords "关键词1,关键词2"]

输出: llm_backend/knowledge_data/product_knowledge_docx/
  每个商品一份 docx: H1=商品名, 章节=属性知识库/商品描述/商品摘要(有数据才生成)

说明:
  - 关键词筛选命中任一即纳入; 默认智能家具关键词, 可用 --keywords 覆盖
  - 按商品名去重
  - 商品名从常见字段自动识别, 识别不到则跳过并计数
"""

import argparse
import json
import re
import sys
from pathlib import Path

from docx import Document

# ── 路径配置 ──────────────────────────────────────────────
PROJECT_ROOT = Path(__file__).resolve().parent.parent
OUTPUT_DIR = PROJECT_ROOT / "llm_backend" / "knowledge_data" / "product_knowledge_docx"

# 智能家具关键词（命中任一即纳入）。京东"家具"归属家电/家居类目，
# 实际数据中智能家具多为 按摩椅/晾衣架/窗帘/升降桌/床垫 等。
DEFAULT_KEYWORDS = (
    "沙发,按摩椅,电动床,智能床,床垫,升降桌,书桌,电脑桌,办公桌,餐桌,茶几,"
    "床头柜,衣柜,电视柜,鞋柜,梳妆台,晾衣架,窗帘,智能锁,智能家具,置物架,书架"
)

# 字段名候选（小写匹配，取第一个命中的非空值）
TITLE_FIELDS = ["title", "product_name", "productname", "name", "sku_name", "sku", "product"]
ATTR_FIELDS = ["attribute", "attributes", "attr", "knowledge", "kb", "attrs", "property"]
DESC_FIELDS = ["input_text", "inputtext", "text", "desc", "description", "content", "input"]
SUMMARY_FIELDS = ["summary", "output", "output_summary", "outputsummary", "gt", "answer", "label"]


def find_field(rec: dict, candidates: list) -> str:
    """按候选字段名取第一个非空字符串值（兼容字符串/list/dict）。"""
    for f in candidates:
        if f in rec:
            v = rec[f]
            if isinstance(v, str):
                return v
            if isinstance(v, list):
                return "\n".join(str(x) for x in v)
            if isinstance(v, dict):
                return "\n".join(f"{k}：{v}" for k, v in v.items())
    return ""


def load_records(data_dir: Path) -> list[dict]:
    """自动探测并加载数据目录下的 jsonl/json/csv/txt 文件。"""
    records: list[dict] = []
    for fp in sorted(data_dir.iterdir()):
        if not fp.is_file():
            continue
        suffix = fp.suffix.lower()
        try:
            if suffix == ".jsonl":
                with open(fp, encoding="utf-8") as f:
                    for line in f:
                        line = line.strip()
                        if line:
                            records.append(json.loads(line))
            elif suffix == ".json":
                data = json.loads(Path(fp).read_text(encoding="utf-8"))
                if isinstance(data, list):
                    records.extend(data)
                elif isinstance(data, dict):
                    for v in data.values():
                        if isinstance(v, list):
                            records.extend(v)
            elif suffix == ".csv":
                import csv

                with open(fp, encoding="utf-8", newline="") as f:
                    records.extend(csv.DictReader(f))
            elif suffix == ".txt":
                # 文本文件: 空行分隔的块, 每块尽力取首行做标题, 其余为描述
                blocks = re.split(r"\n\s*\n", Path(fp).read_text(encoding="utf-8"))
                for b in blocks:
                    lines = [l.strip() for l in b.splitlines() if l.strip()]
                    if lines:
                        records.append({"title": lines[0], "input_text": "\n".join(lines[1:])})
        except Exception as e:
            print(f"[WARN] 解析 {fp.name} 失败, 跳过: {e}")
    return records


def sanitize_filename(name: str, max_len: int = 30) -> str:
    """文件名安全化（去非法字符, 截断）。"""
    name = re.sub(r'[\\/:*?"<>|\s]+', "_", name)
    return name[:max_len].strip("_")


def build_docx(rec: dict, out_path: Path) -> bool:
    """单商品生成 docx。返回是否成功（标题字段缺失视为失败）。"""
    title = find_field(rec, TITLE_FIELDS).strip()
    if not title:
        return False

    attr = find_field(rec, ATTR_FIELDS).strip()
    desc = find_field(rec, DESC_FIELDS).strip()
    summary = find_field(rec, SUMMARY_FIELDS).strip()

    doc = Document()
    doc.add_heading(title, level=1)

    if attr:
        doc.add_heading("商品属性", level=2)
        for line in attr.splitlines():
            line = line.strip()
            if line:
                doc.add_paragraph(line, style="List Bullet")
    if desc:
        doc.add_heading("商品描述", level=2)
        for para in re.split(r"\n\s*\n", desc):
            para = para.strip()
            if para:
                doc.add_paragraph(para)
    if summary:
        doc.add_heading("商品摘要", level=2)
        doc.add_paragraph(summary)

    doc.save(str(out_path))
    return True


def main():
    parser = argparse.ArgumentParser(description="jd-aig CEPSUM 数据 → docx 商品知识文档")
    parser.add_argument("--data-dir", required=True, help="数据集目录（jsonl/json/csv/txt 均支持）")
    parser.add_argument("--limit", type=int, default=50, help="最多生成多少份 docx（默认 50）")
    parser.add_argument("--keywords", default=DEFAULT_KEYWORDS, help="筛选关键词, 逗号分隔")
    args = parser.parse_args()

    data_dir = Path(args.data_dir)
    if not data_dir.is_dir():
        print(f"[ERROR] 数据目录不存在: {data_dir}")
        sys.exit(1)

    keywords = [k.strip() for k in args.keywords.split(",") if k.strip()]
    print(f"加载数据: {data_dir}")
    records = load_records(data_dir)
    print(f"共 {len(records)} 条记录")

    # 关键词筛选 + 按商品名去重（保留先出现的）
    seen, matched = set(), []
    for rec in records:
        title = find_field(rec, TITLE_FIELDS).strip()
        if not title:
            continue
        text = " ".join(find_field(rec, TITLE_FIELDS + DESC_FIELDS + ATTR_FIELDS).split())
        if any(k in text for k in keywords) and title not in seen:
            seen.add(title)
            matched.append(rec)
    print(f"命中 {len(matched)} 条（去重后, 目标 {args.limit}）")

    # 生成 docx
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    created = skipped = 0
    for i, rec in enumerate(matched[: args.limit], start=1):
        title = find_field(rec, TITLE_FIELDS).strip()
        out = OUTPUT_DIR / f"{i:03d}_{sanitize_filename(title)}.docx"
        if build_docx(rec, out):
            created += 1
            print(f"  -> {out.name}")
        else:
            skipped += 1

    print(f"\n完成: 生成 {created} 份 docx, 跳过 {skipped} 条(缺标题), 输出目录: {OUTPUT_DIR}")
    if len(matched) < args.limit:
        print(f"[WARN] 仅命中 {len(matched)} 条, 不足 {args.limit}, 可放宽 --keywords")


if __name__ == "__main__":
    main()
