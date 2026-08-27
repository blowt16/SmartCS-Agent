"""
生成《京东自营售后政策.docx》独立知识文档(与商品知识文档同目录,一并上传索引)。

文档结构(与商品文档同款式体系,样式/召回约束见 spec §3.2/§3.3):
  H1 京东自营售后政策
  ├─ 总述段(Normal): 适用范围
  ├─ H2 章节 ×5(每章 ≥450 字符,含标题与条款,强制独立成块)
  │   └─ List Bullet 条款(高频业务词,提升 BM25 命中)
  └─ H2 参考来源(Normal 链接)

条款来源: 京东帮助中心官方售后政策(2026-08-27 检索, 见 docs/superpowers/notes/2026-08-27-京东自营售后政策检索.md)

用法:
  python scripts/build_jd_aftersales_docx.py
输出: llm_backend/knowledge_data/product_knowledge_docx/京东自营售后政策.docx (gitignore 内)
"""

import argparse
from pathlib import Path

from docx import Document

PROJECT_ROOT = Path(__file__).resolve().parent.parent
OUTPUT_DIR = PROJECT_ROOT / "llm_backend" / "knowledge_data" / "product_knowledge_docx"
DOC_TITLE = "京东自营售后政策"

# 总述段(桥接块: "XX商品+售后主题" query 的召回入口,spec §3.3)
PREAMBLE = (
    "本政策适用于所有京东自营商品（商品信息中标明“京东自营”的商品）。"
    "京东自营商品由京东直接销售并提供售后，以下为售后政策要点整理，"
    "具体以订单售后页面与京东平台最新政策为准。"
)

# 每章节 (标题, [条款,...]); 内容自然长度,不约束字符量(生产环境内容量随机)
POLICY_SECTIONS = [
    ("一、七天无理由退货", [
        "京东自营商品自签收次日起 7 天内支持无理由退货，法律规定或平台明确不可退货的商品除外",
        "完好标准：商品保持原有品质与功能，商品本身、配件、商标标识齐全，不影响二次销售",
        "不支持无理由退货：个人定制类、鲜活易腐类、已拆封数字化商品（音像软件）、防伪码刮开、激活类商品",
        "家具类商品安装后不影响七天无理由退货，但需保持商品完好，包装与配件完整",
        "退货时赠品需一并退回；赠品缺失将影响主商品退款金额",
        "申请路径：订单详情 → 申请售后 → 选择“七天无理由退货” → 填写原因提交",
    ]),
    ("二、价格保护", [
        "京东自营商品支持价格保护：一般商品价保期 7 天，大家电等部分品类价保期 30 天",
        "带“30-30-180”标识的自营商品，签收后 30 天内降价可申请价保返还差价",
        "申请路径：APP 端“我的 → 客户服务 → 价格保护 → 申请价保”，系统自动核算应退差价并原路退回",
        "不支持价保：非京东购买商品、无效订单、已申请售后服务的订单、申请时无货或参与秒杀的商品",
        "价保期内同一商品仅可申请一次价保，以提交时间对应的促销价格核算",
    ]),
    ("三、售后流程", [
        "申请入口：PC 端“客户服务 → 售后服务 → 返修/退换货申请”；APP 端“我的 → 退换/售后 → 申请售后”",
        "流程：申请售后 → 选择退货/换货/维修/价保类型 → 填写原因、上传凭证 → 选择上门取件时间 → 提交",
        "审核：京东受理并审核申请，通常 1 个工作日内完成审核并反馈结果",
        "处理：商品检测后按申请执行退款/换新/维修，退款原路退回",
        "京东自营商品可直接发起售后；第三方商家商品需先与卖家沟通协商后再处理",
    ]),
    ("四、运费与上门取件", [
        "京东自营商品因质量问题退换货，往返运费由京东承担，不收取用户任何运费",
        "七天无理由退货由买家承担寄回运费；钻石级别客户及 PLUS 会员可享免费或双向免费服务",
        "京东上门取件覆盖范围以收货地址为准；超出范围可自行寄回，凭订单号、快递单号及费用凭证返还运费",
        "大件商品（如沙发、床等家具）由京东物流上门取件，无需用户自行搬运",
        "取件时间可在线预约，支持改期；长时间无法取件时客服会电话联系协调",
    ]),
    ("五、商品类目差异说明", [
        "质量问题退换货标准：7 天内退货、15 天内换货，免收返回运费",
        "物流损、缺件或商品描述不符：支持 7 天内退货、15 天内换货，免收运费",
        "个人原因退换货：商品完好前提下支持 7 天内退货，不支持 15 天内换货，返回运费由买家承担",
        "不予办理退换货：过保商品、未经授权维修/人为损坏/进液、无法提供三包凭证或凭证信息不符",
        "带“30-30-180”标识自营商品：30 天质量问题退货、180 天质量问题取回检测换新或上门换新，各享 1 次",
        "商品页面标注的专属售后承诺（如质保年限、免费安装）优先于通用政策",
    ]),
]

SOURCE_SECTIONS = [
    ("京东帮助中心-常见问题分类（售后）", "https://help.jd.com/user/issue/942-3879.html"),
    ("京东自营商品售后", "https://in.m.jd.com/help/app/ziyingshouhou.html"),
    ("京东帮助中心入口", "https://help.jd.com/"),
]


def build_docx(out_path: Path) -> None:
    doc = Document()
    doc.add_heading(DOC_TITLE, level=1)
    doc.add_paragraph(PREAMBLE)

    for title, bullets in POLICY_SECTIONS:
        doc.add_heading(title, level=2)
        for b in bullets:
            doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("参考来源", level=2)
    for name, url in SOURCE_SECTIONS:
        doc.add_paragraph(f"{name}: {url}")

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def main():
    parser = argparse.ArgumentParser(description="生成京东自营售后政策知识文档")
    args = parser.parse_args()
    out_path = OUTPUT_DIR / f"{DOC_TITLE}.docx"
    build_docx(out_path)
    total = sum(len(title) + sum(len(b) + 2 for b in bullets)
                for title, bullets in POLICY_SECTIONS)
    print(f"完成: 生成 {out_path}（{len(POLICY_SECTIONS)} 个章节, 条款总长 {total} 字符）")


if __name__ == "__main__":
    main()
